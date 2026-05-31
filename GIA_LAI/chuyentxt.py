from pathlib import Path
from docx import Document


# =========================
# 1. NHẬP ĐƯỜNG DẪN THƯ MỤC GỐC Ở ĐÂY
# =========================

ROOT_PATH = r"C:\Users\nguye\Downloads\GIA_LAI-20260528T093821Z-3-001\GIA_LAI"

# True = chuyển xong sẽ xóa file .docx
# False = chỉ chuyển, chưa xóa
DELETE_DOCX_AFTER_CONVERT = True


# =========================
# 2. XÓA MẤY DÒNG ĐẦU KHÔNG CẦN
# =========================

def is_metadata_line(text):
    text_lower = text.strip().lower()

    if not text_lower:
        return True

    remove_prefixes = [
        "địa điểm:",
        "định dạng:",
        "mục tiêu:",
    ]

    if text_lower.endswith(".docx") or text_lower == ".docx":
        return True

    for prefix in remove_prefixes:
        if text_lower.startswith(prefix):
            return True

    return False


# =========================
# 3. NHẬN DIỆN TIÊU ĐỀ IN ĐẬM
# =========================

def is_bold_heading(paragraph):
    text = paragraph.text.strip()

    if not text:
        return False

    if is_metadata_line(text):
        return False

    style_name = paragraph.style.name.lower() if paragraph.style else ""

    if "heading" in style_name or "title" in style_name:
        return True

    total_chars = 0
    bold_chars = 0

    for run in paragraph.runs:
        run_text = run.text.strip()

        if not run_text:
            continue

        total_chars += len(run_text)

        if run.bold:
            bold_chars += len(run_text)

    if total_chars == 0:
        return False

    bold_ratio = bold_chars / total_chars

    return bold_ratio >= 0.7


# =========================
# 4. ĐỌC FILE DOCX
# =========================

def read_docx(file_path):
    doc = Document(file_path)
    text_list = []
    section_number = 1

    # Đọc đoạn văn
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Xóa dòng đầu như .docx, Địa điểm, Định dạng, Mục tiêu
        if is_metadata_line(text):
            continue

        # Dòng in đậm thì đánh số mục
        if is_bold_heading(para):
            text_list.append(f"{section_number}. {text}")
            section_number += 1
        else:
            text_list.append(text)

    # Đọc bảng nếu có
    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_text.append(cell_text)

            if any(row_text):
                text_list.append("\t".join(row_text))

    return "\n\n".join(text_list)


# =========================
# 5. CHUYỂN TOÀN BỘ DOCX SANG TXT RỒI XÓA DOCX
# =========================

def convert_all_docx_to_txt_and_delete(root_path):
    root_path = Path(root_path)

    if not root_path.exists():
        print("Đường dẫn không tồn tại.")
        return

    docx_files = [
        file for file in root_path.rglob("*.docx")
        if not file.name.startswith("~$")
    ]

    if not docx_files:
        print("Không tìm thấy file .docx nào.")
        return

    print(f"Tìm thấy {len(docx_files)} file .docx")
    print("-" * 60)

    success_count = 0
    delete_count = 0
    error_count = 0

    for docx_file in docx_files:
        try:
            txt_file = docx_file.with_suffix(".txt")

            text = read_docx(docx_file)

            with open(txt_file, "w", encoding="utf-8") as f:
                f.write(text)

            success_count += 1
            print(f"Đã chuyển: {docx_file} -> {txt_file.name}")

            if DELETE_DOCX_AFTER_CONVERT:
                if txt_file.exists():
                    docx_file.unlink()
                    delete_count += 1
                    print(f"Đã xóa DOCX: {docx_file.name}")

        except Exception as e:
            error_count += 1
            print(f"Lỗi file: {docx_file}")
            print(f"Chi tiết lỗi: {e}")

        print("-" * 60)

    print("HOÀN TẤT")
    print(f"Chuyển thành công: {success_count}")
    print(f"Đã xóa DOCX: {delete_count}")
    print(f"Lỗi: {error_count}")


# =========================
# 6. CHẠY TOOL
# =========================

convert_all_docx_to_txt_and_delete(ROOT_PATH)